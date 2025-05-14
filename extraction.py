from notion_client import Client as NotionClient
from langsmith import Client as LangSmithClient
from langchain.prompts import PromptTemplate
from langchain_core.prompts import ChatPromptTemplate
import openai
import unicodedata
import os
from io import StringIO
import re
from docx import Document
from docx.shared import RGBColor
from typing import Optional
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

class NotionLangSmithSync:

    TAG_STYLES = {
        "ONLY_ERP_VALUE": RGBColor(255, 105, 180),     # Pink
        "FUNCTION_VALUE": RGBColor(30, 144, 255),      # Blue
        "UPDATE_ERP_VALUE": RGBColor(128, 0, 32),      # Dark red
        "FUNCTION_VALUE_BUSINESS_DESCRIPTION": RGBColor(30, 255, 100),  # Green
    }
 
    TAG_PATTERN = re.compile(
        r"(ONLY_ERP_VALUE|FUNCTION_VALUE|UPDATE_ERP_VALUE|FUNCTION_VALUE_BUSINESS_DESCRIPTION)\[(.*?)\]\1",
        re.DOTALL
    )
    PROMPT_TAGS = {
    "mvresolvers": "agent:ChatGPT MV Resolvers",
    "ccresolvers": "agent:ChatGPT CC Resolvers",
    "mvsales": "agent:MV Sales Agent",
    "ccsales": "agent:CC Sales Agent",
    "delighters": "agent:Delighters Agent",
    "doctors": "agent:Doctor's assistant",
    "maidsat": "agent:Maids Line",
    "atafrican": "agent:Maids Line",
    "atphilipina": "agent:Maids Line",
    "atethiopian": "agent:Maids Line",
}

    def __init__(self, erp_value_option = 0, function_value_option = 0, update_erp_value_option = 4, langsmith_api_key: str = None, notion_database_id: str = None):
        self.notion_token = os.getenv("NOTION_TOKEN")
        self.langsmith_api_key = langsmith_api_key
        self.notion = NotionClient(auth=self.notion_token)
        self.langsmith = LangSmithClient(api_key=self.langsmith_api_key)
        self.NOTION_DATABASE_ID = notion_database_id
        self.OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
        self.erp_value_option = erp_value_option
        self.function_value_option = function_value_option
        self.update_erp_value_option = update_erp_value_option

    def get_all_blocks(self, block_id, level=0):
        blocks = []
        children = self.notion.blocks.children.list(block_id)
        for block in children["results"]:
            block["level"] = level
            blocks.append(block)
            if block.get("has_children", False):
                blocks.extend(self.get_all_blocks(block["id"], level + 1))
        return blocks

    def get_page_title(self, page_id: str) -> str:
        try:
            page = self.notion.pages.retrieve(page_id=page_id)
            for prop in page["properties"].values():
                if prop["type"] == "title":
                    raw_title = "".join(part["text"]["content"] for part in prop["title"])
                    sanitized = unicodedata.normalize("NFKD", raw_title).encode("ascii", "ignore").decode("ascii")
                    sanitized = sanitized.lower()
                    sanitized = re.sub(r'[^a-z0-9_-]', '_', sanitized)
                    sanitized = re.sub(r'^[^a-z]+', '', sanitized)
                    return sanitized or "default_prompt"
            return "untitled_prompt"
        except Exception as e:
            print(f"Error retrieving or sanitizing title for page {page_id}: {e}")
            return "error_prompt"

    def get_notion_variable_value(self, database_id, variable_name):
        response = self.notion.databases.query(database_id=database_id)
        value = ""
        for row in response["results"]:
            props = row["properties"]
            name_prop = props.get("Name", {}).get("title", [])
            name = name_prop[0].get("plain_text", "").strip() if name_prop else ""
            default_value_prop = props.get("Default Value", {}).get("rich_text", [])
            default_value = default_value_prop[0].get("plain_text", "").strip() if default_value_prop else ""
            if name == variable_name:
                value = default_value
        return value

    def get_function_value_business_description(self, text):
      client = openai.OpenAI(api_key=self.OPENAI_API_KEY)  # Replace with your real key or use environment variable

      system_prompt = """
        # FUNCTION_VALUE MEANING:
        The FUNCTION_VALUE block contains structured logic or condition-based formulas related to ERP or business processes. It defines how the system should behave in specific functional scenarios.

        # TASK:
        Summarize the meaning and business purpose of the FUNCTION_VALUE logic into a concise two-sentence description.

        # INSTRUCTIONS:
        - Your output must be exactly two sentences.
        - Rephrase complex conditions into business-friendly terms.
        - Describe what the logic is doing and why it matters from a business perspective.
        - Always start your description by mentioning the FUNCTION_VALUE.
        - Do NOT include code or pseudo-code.
        - Write in clear, plain English for non-technical stakeholders.

        # INPUT FORMAT:
        The input contains the full FUNCTION_VALUE content including logic, nested conditions, and references to ERP variables.
        """

      response = client.chat.completions.create(
          model="gpt-4o",
          temperature=0.2,
          messages=[
              {"role": "system", "content": system_prompt.strip()},
              {"role": "user", "content": text.strip()}
          ]
      )

      return response.choices[0].message.content.strip()

    def get_update_erp_value_business_description(self, text):
      client = openai.OpenAI(api_key=self.OPENAI_API_KEY)

      system_prompt = """
        # UPDATE_ERP_VALUE MEANING:
        The UPDATE_ERP_VALUE block contains logic related to updating or modifying ERP values based on certain conditions. It defines how the system should modify or update data in the ERP system under specific circumstances.

        # TASK:
        Summarize the meaning and business purpose of the UPDATE_ERP_VALUE logic into a concise two-sentence description.

        # INSTRUCTIONS:
        - Your output must be exactly two sentences.
        - Rephrase complex conditions into business-friendly terms.
        - Describe what the logic is doing and why it matters from a business perspective.
        - Always start your description by mentioning the UPDATE_ERP_VALUE.
        - Do NOT include code or pseudo-code.
        - Write in clear, plain English for non-technical stakeholders.

        # INPUT FORMAT:
        The input contains the full UPDATE_ERP_VALUE content including logic, nested conditions, actions to be triggered and references to ERP variables.
        """

      response = client.chat.completions.create(
          model="gpt-4o",
          temperature=0.2,
          messages=[
              {"role": "system", "content": system_prompt.strip()},
              {"role": "user", "content": text.strip()}
          ]
      )

      return response.choices[0].message.content.strip()

    def fetch_all_block_content(self, page_id, block_id, add_title = True):
      """
      Fetches all child blocks under a given block_id, including nested blocks.
      Formats each block as a list of strings with proper indentation and hierarchical numbering.
      """
      content_lines = []
      ignored_blocks = []
      list_counters = {}
      # Fetch the block itself
      block = self.notion.blocks.retrieve(block_id)
      block_type = block["type"]
      block_data = block.get(block_type, {})
      block_text = self.extract_text(block_data)

      if add_title and block_text:
          indent = "    " * 1
          if block_type == "bulleted_list_item":
              content_lines.append(f"{indent}- {block_text}")
          elif block_type == "numbered_list_item":
              list_counters[1] = list_counters.get(1, 0) + 1
              content_lines.append(f"{indent}{list_counters[1]}. {block_text}")
          else:
              content_lines.append(f"{indent}{block_text}")

      # Fetch all nested child blocks
      all_nested_blocks = self.get_all_blocks(block_id)

      for block in all_nested_blocks:
          block_type = block["type"]
          block_data = block.get(block_type, {})
          if block["parent"].get("block_id") in ignored_blocks:
              ignored_blocks.append(block["id"])
              continue

          level = block.get("level", 0)
          indent = "    " * (level + 1)

          if block_type == "toggle":
              text = self.extract_text(block["toggle"])
              text_lower = text.lower()

              if "function_value" in text_lower:
                  ignored_blocks.append(block["id"])
                  formatted_function_value = self.get_function_value(page_id, block["id"])
                  content_lines.append(f"{indent}{formatted_function_value}")
              elif "update_erp_value" in text_lower:
                  ignored_blocks.append(block["id"])
                  formatted_update_erp_value = self.get_update_erp_value(page_id, block["id"])
                  content_lines.append(f"{indent}{formatted_update_erp_value}")
              else:
                  if text:
                      content_lines.append(f"{indent}{text}")
          else:
              if not block_data:
                  continue

              block_text = self.extract_text(block_data)
              if not block_text:
                  continue

              if block_type == "bulleted_list_item":
                  content_lines.append(f"{indent}- {block_text}")
              elif block_type == "numbered_list_item":
                  # Ensure counter for this level exists
                  while len(list_counters) <= level:
                      list_counters[len(list_counters)] = 0
                  list_counters[level] += 1

                  # Reset deeper levels
                  for deeper_level in list(filter(lambda l: l > level, list_counters)):
                      list_counters[deeper_level] = 0

                  hierarchy = [str(list_counters[l]) for l in sorted(list_counters) if l <= level and list_counters[l] > 0]
                  content_lines.append(f"{indent}{'.'.join(hierarchy)}. {block_text}")
              else:
                  content_lines.append(f"{indent}{block_text}")

      if not add_title:
        content_lines.pop(0)
      return "\n".join(content_lines)


    def get_erp_value(self, page_id):
        try:
            page = self.notion.pages.retrieve(page_id=page_id)
            val = "Untitled Page"
            title_prop = page["properties"].get("Name")
            if title_prop and title_prop.get("type") == "title":
              title_parts = title_prop["title"]
              val =  "".join(part["text"]["content"] for part in title_parts)
            if self.erp_value_option == 0:
              return "ONLY_ERP_VALUE[" + self.get_notion_variable_value(self.NOTION_DATABASE_ID, val) + "]ONLY_ERP_VALUE"
            if self.erp_value_option == 2:
              return "ONLY_ERP_VALUE[" + f"{{{val}}}" + "]ONLY_ERP_VALUE"
            return "ONLY_ERP_VALUE[" + f"@{val}@" + "]ONLY_ERP_VALUE"

        except Exception as e:
            print(f"Error retrieving page title for {page_id}: {e}")
            return "ONLY_ERP_VALUE[" + "@Error@" + "]ONLY_ERP_VALUE"

    def get_function_value(self, page_id, block_id):
        try:
            page = self.notion.pages.retrieve(page_id=page_id)
            val = "@Function Name@"
            # If the function_value_option is 3, fetch the business description
            if self.function_value_option == 3:
                  val = self.get_function_value_business_description(self.fetch_all_block_content(page_id, block_id))
                  return "FUNCTION_VALUE_BUSINESS_DESCRIPTION[" + val + "]FUNCTION_VALUE_BUSINESS_DESCRIPTION"  # Return the extracted value

            # If the function_value_option is 4, fetch the whole content as it is
            elif self.function_value_option == 4:
                  val = self.fetch_all_block_content(page_id, block_id)
                  return "FUNCTION_VALUE[" + val + "]FUNCTION_VALUE"

            elif self.function_value_option == 2:
                blocks = self.get_all_blocks(block_id)
                for block in blocks:
                    if block['type'] == 'callout' and block.get('has_children'):
                        val = self.fetch_all_block_content(page_id, block["id"], False)
                        return "FUNCTION_VALUE[" + val + "]FUNCTION_VALUE"

            else:
              blocks = self.get_all_blocks(block_id)
              for block in blocks:
                  if block['type'] == 'paragraph' and 'rich_text' in block['paragraph']:
                      block_content = self.extract_text(block['paragraph'])

                      if self.function_value_option == 0 and block_content.startswith("Value if no condition is met:"):
                            val = block_content.split("Value if no condition is met:")[1].strip()
                            return "FUNCTION_VALUE[" + val + "]FUNCTION_VALUE"  # Return the extracted value

                      elif self.function_value_option == 1 and block_content.startswith("Name*:"):
                            val = block_content.split("Name*:")[1].strip()
                            return "FUNCTION_VALUE[" + f"@{val}@" + "]FUNCTION_VALUE"  # Return the extracted value

            return "FUNCTION_VALUE[" + val + "]FUNCTION_VALUE"
        except Exception as e:
            print(f"Error retrieving function value for {page_id}: {e}")
            return "FUNCTION_VALUE[" + "@Error@" + "]FUNCTION_VALUE"

    def get_update_erp_value(self, page_id, block_id):
      try:
          page = self.notion.pages.retrieve(page_id=page_id)
          val = "@Update ERP Value Name@"

           # If the update_erp_value_option is 2, fetch the business description
          if self.update_erp_value_option == 2:
                val = self.get_update_erp_value_business_description(self.fetch_all_block_content(page_id, block_id))
                return "UPDATE_ERP_VALUE[" + val + "]UPDATE_ERP_VALUE"  # Return the extracted value

          # If the update_erp_value_option is 3, fetch the whole content as it is
          elif self.update_erp_value_option == 3:
                val = self.fetch_all_block_content(page_id, block_id)
                return "UPDATE_ERP_VALUE[" + val + "]UPDATE_ERP_VALUE"  # Return the extracted value
          # If the update_erp_value_option is 4, return nothing
          elif self.update_erp_value_option == 4:
                return ""
          else:

            blocks = self.get_all_blocks(block_id)

            # Iterate through the blocks and check the condition based on the function_value_option
            for block in blocks:
                if block['type'] == 'paragraph' and 'rich_text' in block['paragraph']:
                    block_content = self.extract_text(block['paragraph'])

                    # If the update_erp_value_option is 0, look for the content after "Value if no condition is met:"
                    if self.update_erp_value_option == 0:
                        if block_content.startswith("Value if no condition is met:"):
                            val = block_content.split("Value if no condition is met:")[1].strip()
                            return "UPDATE_ERP_VALUE[" + val + "]UPDATE_ERP_VALUE"

                    # If the update_erp_value_option is 1, look for the content after "Name*:"
                    elif self.update_erp_value_option == 1:
                        if block_content.startswith("Name*:"):
                            val = block_content.split("Name*:")[1].strip()
                            return "UPDATE_ERP_VALUE[" + f"@{val}@" + "]UPDATE_ERP_VALUE"

            return "UPDATE_ERP_VALUE[" + val + "]UPDATE_ERP_VALUE"

      except Exception as e:
          print(f"Error retrieving update erp value for {page_id}: {e}")
          return "UPDATE_ERP_VALUE[" + "@Error@" + "]UPDATE_ERP_VALUE"

    def extract_text(self, block_content):
        try:
            parts = []
            for rt in block_content["rich_text"]:
                if rt.get("type") == "mention":
                    mention = rt.get("mention", {})
                    if mention.get("type") == "page":
                        page_id = mention["page"].get("id")
                        if page_id:
                            formatted_erp_value = self.get_erp_value(page_id)
                            parts.append(formatted_erp_value)
                        else:
                            parts.append("{Unknown Page}")
                    else:
                        parts.append("{Mention}")
                elif rt.get("type") == "text":
                    parts.append(rt["text"].get("content", ""))
            return "".join(parts)
        except (KeyError, IndexError, TypeError):
            return ""

    def get_notion_prompt(self, page_id):
        all_blocks = self.get_all_blocks(page_id)
        prompt = []
        list_counters = {}
        ignored_blocks = []

        for block in all_blocks:
            block_type = block["type"]
            level = block.get("level", 0)
            indent = "    " * level

            if block["parent"].get("block_id") in ignored_blocks:
                ignored_blocks.append(block["id"])
                continue

            if block_type == "toggle":
                text = self.extract_text(block["toggle"])
                text_lower = text.lower()

                if "function_value" in text_lower:
                    ignored_blocks.append(block["id"])
                    formatted_function_value = self.get_function_value(page_id, block["id"])
                    prompt.append(f"{indent}{formatted_function_value}")
                elif "update_erp_value" in text_lower:
                    ignored_blocks.append(block["id"])
                    formatted_update_erp_value = self.get_update_erp_value(page_id, block["id"])
                    prompt.append(f"{indent}{formatted_update_erp_value}")
                continue

            block_data = block.get(block_type)
            if not block_data:
                continue

            text = self.extract_text(block_data)
            if not text:
                continue

            if block_type == "bulleted_list_item":
                formatted = f"{indent}- {text}"
            elif block_type == "numbered_list_item":
                # Ensure counters exist up to current level
                while len(list_counters) <= level:
                    list_counters[len(list_counters)] = 0
                list_counters[level] += 1

                # Reset all deeper levels
                for deeper_level in list(filter(lambda l: l > level, list_counters)):
                    list_counters[deeper_level] = 0

                # Build hierarchical number like 1.2.3
                hierarchy = [str(list_counters[l]) for l in sorted(list_counters) if l <= level and list_counters[l] > 0]
                formatted = f"{indent}{'.'.join(hierarchy)}. {text}"
            else:
                formatted = f"{indent}{text}"

            prompt.append(formatted)

        return prompt

    def strip_custom_identifiers(self, text):
      # Removes FUNCTION_VALUE[], ONLY_ERP_VALUE[], UPDATE_ERP_VALUE[] identifiers, keeping only content inside
      text = re.sub(r"FUNCTION_VALUE(?:_BUSINESS_DESCRIPTION)?\[(.*?)\]FUNCTION_VALUE(?:_BUSINESS_DESCRIPTION)?", r"\1", text, flags=re.DOTALL)
      text = re.sub(r"UPDATE_ERP_VALUE\[(.*?)\]UPDATE_ERP_VALUE", r"\1", text, flags=re.DOTALL)
      text = re.sub(r"ONLY_ERP_VALUE\[(.*?)\]ONLY_ERP_VALUE", r"\1", text, flags=re.DOTALL)
      return text

    def get_all_prompt_names_and_ids(self):
        """
        Fetches all prompt names and their UUIDs from LangSmith.
        Returns:
            A dictionary: {prompt_name: prompt_id}
        """
        prompt_map = {}
        response = self.langsmith.list_prompts(is_public=False, limit=100)

        # LangSmith Client returns paginated generator
        for prompt in response.repos:
            print(prompt)
            name = prompt.repo_handle
            id = str(prompt.id)
            prompt_map[name] = id

        return prompt_map

    def sync_prompt(self, page_id: str, export: bool, prompt_name: str, department = None):
        paragraphs = self.get_notion_prompt(page_id)
        system_message = self.strip_custom_identifiers("\n".join(paragraphs))

        if not export:
            # Save system_message to a text file
            filename = f"{prompt_name or 'notion_prompt'}.txt"
            file_path = Path("outputs") / filename
            os.makedirs(file_path.parent, exist_ok=True)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(system_message)
            return str(file_path)

        chat_prompt_template = ChatPromptTemplate([
            ("system", system_message),
            ("user", "{Conversation}"),
        ])

        prompt_identifier = prompt_name.strip().replace(" ", "_").lower()

        print(f"\nUsing prompt identifier: {prompt_identifier}")
        print(system_message)
        print(f"\nTotal lines in prompt: {len(paragraphs)}")

        try:
            essential_tags = ["notion"]
            if department and department in self.PROMPT_TAGS:
                essential_tags.append(self.PROMPT_TAGS[department])
            else:
                essential_tags.append("agent:None")
            tags = []
            if self.langsmith._prompt_exists(prompt_identifier):
                tags = self.langsmith.get_prompt(prompt_identifier).tags
            for tag in essential_tags:
                if tag not in tags:
                    tags.append(tag)

            prompt_url = self.langsmith.push_prompt(
                prompt_identifier=prompt_identifier,
                object=chat_prompt_template,
                description="Updated prompt with content",
                tags=tags
            )
            print(f" Successfully pushed prompt: {prompt_identifier}")
            return prompt_url

        except Exception as e:
            error_msg = str(e)
            if "Nothing to commit" in error_msg:
                print(f" No changes in prompt: {prompt_identifier}. Skipping push.")
                return "Prompt not updated — no changes detected."
            else:
                raise e

    def process_text(self, paragraph, text):
      pos = 0
      for match in self.TAG_PATTERN.finditer(text):
          start, end = match.span()
          tag, content = match.groups()

          if start > pos:
              paragraph.add_run(text[pos:start])

          inner_text = StringIO(content).getvalue()
          sub_pos = 0
          for sub_match in self.TAG_PATTERN.finditer(inner_text):
              s_start, s_end = sub_match.span()
              s_tag, s_content = sub_match.groups()

              if s_start > sub_pos:
                  before = inner_text[sub_pos:s_start]
                  run = paragraph.add_run(before)
                  run.font.color.rgb = self.TAG_STYLES[tag]

              run = paragraph.add_run(s_content)
              run.font.color.rgb = self.TAG_STYLES[s_tag]
              sub_pos = s_end

          if sub_pos < len(inner_text):
              tail = inner_text[sub_pos:]
              run = paragraph.add_run(tail)
              run.font.color.rgb = self.TAG_STYLES[tag]

          pos = end

      if pos < len(text):
          paragraph.add_run(text[pos:])

    def add_colored_prompt_to_doc(self, page_id, output_file="outputs/prompt_output.docx"):
        doc = Document()
        prompt_lines = self.get_notion_prompt(page_id)
        for line in prompt_lines:
            paragraph = doc.add_paragraph()
            self.process_text(paragraph, line)
        doc.save(output_file)
        print(f"Document saved to {output_file}")
        return output_file


